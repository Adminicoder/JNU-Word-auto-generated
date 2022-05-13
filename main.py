from docx import Document
from docx.shared import Pt
import random
import datetime
import os

if __name__ == '__main__':
    year = datetime.datetime.today().year
    month = datetime.datetime.today().month
    sum_dict = {}
    name_list = []
    work_path = os.listdir()
    for path in work_path:
        if path.endswith('txt'):
            with open(path, "r", encoding='utf-8') as f:
                data = f.readlines()
                for i in range(0, 3):
                    data[i] = data[i].strip('\n')
                print(data)

            key_wd = ['整理书籍', '打扫卫生', '实验室开关门', '整理资料']
            situation = r'很好/好/一般/差'

            #  工时构建，对于小于16工时的情况不适用
            sum_hrs = float(data[3])  # 总工时
            hr4 = 4.0  # 4小时的数量
            hr2 = (sum_hrs - hr4 * 4) // 2
            hr_rest = (sum_hrs - hr4 * 4) % 2
            if hr2 >= 1:
                hr3 = hr_rest + 2  # 2-4小时
                hr2 = hr2 - 1  # 2小时的数量
            else:
                hr3 = hr_rest
            print('hour_4:', hr4)
            print('hour_2:', hr2)
            print('hour_rest:', hr3)
            print('hour_sum', hr4 * 4 + hr2 * 2 + hr3)

            #  工时随机分配
            hr_list = []
            #  区分3种工时，随机分配，记得hr3是工时大小，不是3小时工时的数量
            #  我找到了一种绝妙的方法解决随机分配问题
            for i in range(0, int(hr4)):
                hr_list.append(4)
            for i in range(0, int(hr2)):
                hr_list.append(2)
            hr_list.append(hr3)
            random.shuffle(hr_list)
            print('hr_list:', hr_list)

            #  日期限定在1-25号
            day = len(hr_list)
            day_list = random.sample(range(1, 25), day)
            print('day_list:', day_list)
            day_list.sort()
            print('sorted_day_list:', day_list)

            #  时间段构建
            full_hour = ':00'
            half_hour = ':30'
            hr2_period = ['8:30-10:30', '10:30-12:30', '14:00-16:00', '18:30-20:30']
            hr4_period = ['8:30-12:30', '14:00-18:00']
            hr3_period = {0.5: '11:30-12:00', 1: '15:00-16:00', 1.5: '14:30-16:00', 2: '18:00-20:00',
                          2.5: '18:30-21:00',
                          3: '9:00-12:00', 3.5: '14:00-17:30'}
            period_list = []
            for i in hr_list:
                if i == 2:
                    ran_i = random.randint(0, 3)
                    period_list.append(hr2_period[ran_i])
                elif i == 4:
                    ran_i = random.randint(0, 1)
                    period_list.append(hr4_period[ran_i])
                else:
                    period_list.append(hr3_period[i])
            print('period_list:', period_list)



            doc = Document('module.docx')
            table = doc.tables
            table[0].cell(0, 6).paragraphs[0].add_run(str(year) + "年" + str(month) + "月").font.size = Pt(12)  # 年月
            table[0].cell(1, 1).paragraphs[0].add_run(data[0]).font.size = Pt(12)  # 姓名
            table[0].cell(1, 6).paragraphs[0].add_run(data[1]).font.size = Pt(12)  # 电话
            table[0].cell(2, 2).paragraphs[0].add_run(data[2]).font.size = Pt(12)  # 银行卡号

            num = len(hr_list)
            for i in range(0, num):
                table[0].cell(i + 4, 0).paragraphs[0].add_run(str(month) + r"." + str(day_list[i])).font.size = Pt(
                    12)  # 日期
                table[0].cell(i + 4, 1).paragraphs[0].add_run(period_list[i]).font.size = Pt(12)
                table[0].cell(i + 4, 3).paragraphs[0].add_run(str(hr_list[i])).font.size = Pt(12)
                ran_context = random.sample(range(0, 3), 2)
                table[0].cell(i + 4, 4).paragraphs[0].add_run(
                    key_wd[ran_context[0]] + r'、' + key_wd[ran_context[1]]).font.size = Pt(12)
                table[0].cell(i + 4, 6).paragraphs[0].add_run(situation).font.size = Pt(12)
            table[0].cell(17, 0).paragraphs[0].add_run(u'计划分配工时：   ' + str(data[3]) + '小时').font.size = Pt(12)
            table[0].cell(17, 0).paragraphs[1].add_run(u'总工时：         ' + str(data[3]) + '小时').font.size = Pt(12)
            doc.save(data[0] + str(month) + '月工作考核表.docx')
            print('Generated\n')
            #  sum_dict = {'龙天翔':{'hr1':2, 'hr_list':[]}}
            sum_dict[data[0]] = {'name': data[0], 'date': day_list, 'period': period_list}
            name_list.append(data[0])
    # print(sum_dict)
    # print(name_list)
    date_list = []
    final_period_list = []
    final_name_list = []
    for i in name_list:
        date_list += sum_dict[i]['date']
        final_period_list += sum_dict[i]['period']
        for j in range(0, len(sum_dict[i]['date'])):
            final_name_list.append(i)
    # print(date_list)
    # print(final_period_list)
    # print(final_name_list)
    n = len(date_list)

    # 遍历所有数组元素
    for i in range(n):
        for j in range(0, n - i - 1):
            if date_list[j] > date_list[j + 1]:
                date_list[j], date_list[j + 1] = date_list[j + 1], date_list[j]
                final_period_list[j], final_period_list[j + 1] = final_period_list[j + 1], final_period_list[j]
                final_name_list[j], final_name_list[j + 1] = final_name_list[j + 1], final_name_list[j]
            elif date_list[j] == date_list[j + 1]:
                if int(final_period_list[j].split('-')[0].split(':')[0]) > int(final_period_list[j+1].split('-')[0].split(':')[0]) :
                    date_list[j], date_list[j + 1] = date_list[j + 1], date_list[j]
                    final_period_list[j], final_period_list[j + 1] = final_period_list[j + 1], final_period_list[j]
                    final_name_list[j], final_name_list[j + 1] = final_name_list[j + 1], final_name_list[j]
                elif final_period_list[j].split('-')[0].split(':')[0] == final_period_list[j+1].split('-')[0].split(':')[0] :
                    if int(final_period_list[j].split('-')[0].split(':')[1]) > int(final_period_list[j+1].split('-')[0].split(':')[1]) :
                        date_list[j], date_list[j + 1] = date_list[j + 1], date_list[j]
                        final_period_list[j], final_period_list[j + 1] = final_period_list[j + 1], final_period_list[j]
                        final_name_list[j], final_name_list[j + 1] = final_name_list[j + 1], final_name_list[j]
    # print(date_list)
    # print(final_period_list)

    doc2 = Document('module2.docx')
    table = doc2.tables

    for i in range(0, len(final_name_list)):
        table[0].cell(i + 2, 0).paragraphs[0].add_run(final_name_list[i])
        table[0].cell(i + 2, 1).paragraphs[0].add_run(str(month) + r"." + str(date_list[i]))
        table[0].cell(i + 2, 2).paragraphs[0].add_run(final_period_list[i].split('-')[0])
        table[0].cell(i + 2, 3).paragraphs[0].add_run(final_period_list[i].split('-')[1])
        table[0].cell(i + 2, 4).paragraphs[0].add_run(final_period_list[i])
    doc2.save('总表.docx')